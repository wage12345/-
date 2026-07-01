import argparse
import importlib.util
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path.cwd()
SKILL_ROOT = Path(
    "C:/Users/塬森/.codex/plugins/cache/openai-primary-runtime/documents/26.623.12021/skills/documents"
)
TABLE_GEOMETRY_PATH = SKILL_ROOT / "scripts" / "table_geometry.py"
LOGO_PATH = Path("C:/Temp/codex-report/media/image1.png")
FONT_SONG = "SimSun"
FONT_HEI = "SimHei"
FONT_ASCII = "Times New Roman"
BODY_SIZE = Pt(12)
CAPTION_SIZE = Pt(10.5)
BODY_LINE_PT = 20
ASSET_DIR = ROOT / "tmp" / "report_assets"
OUTPUT_DIR = ROOT / "deliverables"


def load_table_geometry():
    spec = importlib.util.spec_from_file_location("table_geometry", TABLE_GEOMETRY_PATH)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


table_geometry = load_table_geometry()


def get_font(size, bold=False):
    candidates = [
        ("C:/Windows/Fonts/simhei.ttf", FONT_HEI),
        ("C:/Windows/Fonts/simsun.ttc", FONT_SONG),
        ("C:/Windows/Fonts/msyh.ttc", "Microsoft YaHei"),
    ]
    for path, _name in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def ensure_dirs():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def set_run_font(run, size, bold=False, east_asia=FONT_SONG, ascii_font=FONT_ASCII):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_format(paragraph, *, before=0, after=0, line=BODY_LINE_PT, first_indent=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(line)
    if first_indent is not None:
        fmt.first_line_indent = Cm(first_indent)


def write_paragraph(paragraph, text, *, size=BODY_SIZE, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                    before=0, after=0, line=BODY_LINE_PT, first_indent=None, east_asia=FONT_SONG):
    paragraph.alignment = align
    set_paragraph_format(paragraph, before=before, after=after, line=line, first_indent=first_indent)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, east_asia=east_asia)
    return paragraph


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    write_paragraph(p, text, size=BODY_SIZE, before=0, after=0, line=BODY_LINE_PT, first_indent=0.74)
    return p


def set_outline_level(paragraph, level):
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    size = {1: Pt(16), 2: Pt(14), 3: Pt(12)}[level]
    set_run_font(run, size=size, bold=True)
    set_outline_level(p, level - 1)
    if level == 1:
        set_paragraph_format(p, before=20, after=20, line=BODY_LINE_PT)
    elif level == 2:
        set_paragraph_format(p, before=10, after=10, line=BODY_LINE_PT)
    else:
        set_paragraph_format(p, before=6, after=6, line=BODY_LINE_PT)
    return p


def set_style(style, *, east_asia=FONT_SONG, ascii_font=FONT_ASCII, size=BODY_SIZE, bold=False,
              before=0, after=0, line=BODY_LINE_PT):
    style.font.name = ascii_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    style.font.size = size
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(line)


def configure_styles(doc):
    set_style(doc.styles["Normal"], size=BODY_SIZE, bold=False, before=0, after=0, line=BODY_LINE_PT)
    set_style(doc.styles["Heading 1"], size=Pt(16), bold=True, before=20, after=20, line=BODY_LINE_PT)
    set_style(doc.styles["Heading 2"], size=Pt(14), bold=True, before=10, after=10, line=BODY_LINE_PT)
    set_style(doc.styles["Heading 3"], size=Pt(12), bold=True, before=6, after=6, line=BODY_LINE_PT)


def set_page_layout(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(3.175)
    section.right_margin = Cm(3.175)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
            )


def add_field(paragraph, instruction):
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    run_text = paragraph.add_run(" ")
    set_run_font(run_text, size=Pt(11))

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def restart_page_numbering(section, start=1):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_footer_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run1 = p.add_run("- ")
    set_run_font(run1, size=Pt(11))
    add_field(p, "PAGE")
    run2 = p.add_run(" -")
    set_run_font(run2, size=Pt(11))


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    write_paragraph(p, text, size=CAPTION_SIZE, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                    before=6, after=6, line=16)


def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    write_paragraph(p, text, size=CAPTION_SIZE, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                    before=4, after=10, line=16)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def fill_table_cell(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=BODY_SIZE):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_format(p, before=0, after=0, line=18)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_data_table(doc, headers, rows, widths_cm, title):
    add_table_caption(doc, title)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        fill_table_cell(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, "D9E2F3")

    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 14 else WD_ALIGN_PARAGRAPH.LEFT
            fill_table_cell(table.rows[row_idx].cells[col_idx], str(value), align=align)

    widths_dxa = [int(round(cm / 2.54 * 1440)) for cm in widths_cm]
    table_geometry.apply_table_geometry(
        table,
        table_geometry.exact_column_widths(widths_dxa, total_width_dxa=9360),
        table_width_dxa=9360,
        indent_dxa=120,
    )
    doc.add_paragraph()
    return table


def draw_box(draw, xy, text, font, *, fill="#F5F7FA", outline="#333333", text_fill="#111111"):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    line_heights = []
    max_width = 0
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]
        line_heights.append(height)
        max_width = max(max_width, width)
    total_height = sum(line_heights) + (len(lines) - 1) * 10
    y = y1 + (y2 - y1 - total_height) / 2
    for line, height in zip(lines, line_heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=font, fill=text_fill)
        y += height + 10


def draw_arrow(draw, start, end, fill="#333333"):
    draw.line([start, end], fill=fill, width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 18
    arrow_angle = math.pi / 8
    p1 = (
        end[0] - arrow_len * math.cos(angle - arrow_angle),
        end[1] - arrow_len * math.sin(angle - arrow_angle),
    )
    p2 = (
        end[0] - arrow_len * math.cos(angle + arrow_angle),
        end[1] - arrow_len * math.sin(angle + arrow_angle),
    )
    draw.polygon([end, p1, p2], fill=fill)


def generate_diagrams():
    ensure_dirs()
    title_font = get_font(38)
    box_font = get_font(30)
    small_font = get_font(24)
    assets = {}

    # Figure 3.1
    img = Image.new("RGB", (1800, 1100), "white")
    draw = ImageDraw.Draw(img)
    draw_box(draw, (700, 60, 1100, 180), "用户与浏览器", title_font, fill="#E8F1FB")
    draw_box(draw, (620, 270, 1180, 410), "表示层\nEJS + Bootstrap + 自定义 CSS", box_font)
    draw_box(draw, (220, 520, 620, 680), "控制器层\nAuth / Student / Class", box_font, fill="#F8F5FF")
    draw_box(draw, (700, 520, 1100, 680), "业务模型层\nUser / Student / Class", box_font, fill="#F8F5FF")
    draw_box(draw, (1180, 520, 1580, 680), "工程支撑层\nCodex / Jest / CI / Docker", box_font, fill="#FFF7E8")
    draw_box(draw, (620, 810, 1180, 960), "数据层\nMySQL + schemaSync 初始化脚本", box_font, fill="#EEF7EE")
    draw_arrow(draw, (900, 180), (900, 270))
    draw_arrow(draw, (900, 410), (420, 520))
    draw_arrow(draw, (900, 410), (900, 520))
    draw_arrow(draw, (900, 410), (1380, 520))
    draw_arrow(draw, (420, 680), (860, 810))
    draw_arrow(draw, (900, 680), (900, 810))
    draw_arrow(draw, (1380, 680), (940, 810))
    path = ASSET_DIR / "figure_3_1_architecture.png"
    img.save(path)
    assets["architecture"] = path

    # Figure 3.2
    img = Image.new("RGB", (1800, 1100), "white")
    draw = ImageDraw.Draw(img)
    draw_box(draw, (650, 80, 1150, 220), "学生信息管理系统", title_font, fill="#E8F1FB")
    nodes = [
        ((120, 360, 500, 520), "登录认证\n用户登录 / 会话控制"),
        ((560, 360, 940, 520), "学生管理\n新增 / 编辑 / 删除 / 查询"),
        ((1000, 360, 1380, 520), "班级管理\n班级维护 / 人数统计"),
        ((1440, 360, 1760, 520), "数据导出\nCSV 导出"),
        ((420, 700, 820, 860), "统计展示\n分页 / 状态统计 / 最近录入"),
        ((960, 700, 1380, 860), "工程化能力\n测试 / CI / Docker"),
    ]
    for xy, text in nodes:
        draw_box(draw, xy, text, box_font)
        center = ((xy[0] + xy[2]) / 2, xy[1])
        draw_arrow(draw, (900, 220), center)
    path = ASSET_DIR / "figure_3_2_function.png"
    img.save(path)
    assets["function"] = path

    # Figure 3.3
    img = Image.new("RGB", (1800, 1000), "white")
    draw = ImageDraw.Draw(img)
    draw_box(draw, (150, 300, 550, 620), "users\nid\nusername\npassword", box_font, fill="#F8F5FF")
    draw_box(draw, (700, 180, 1120, 700), "classes\nid\nclass_name\ngrade\nmajor\ncounselor\nmonitor_name\nstudent_capacity\ncreated_at", box_font, fill="#EEF7EE")
    draw_box(draw, (1270, 180, 1670, 820), "students\nid\nstudent_no\nname\ngender\nclass_name\nphone\nemail\nbirthday\nstatus\ndormitory\nemergency_contact\nemergency_phone\nnotes\ncreated_at", box_font, fill="#E8F1FB")
    draw_arrow(draw, (1120, 440), (1270, 440))
    draw.text((1160, 390), "一对多", font=small_font, fill="#111111")
    draw.text((160, 650), "登录账号与密码校验", font=small_font, fill="#333333")
    draw.text((1270, 850), "students.class_name 与 classes.class_name 关联", font=small_font, fill="#333333")
    path = ASSET_DIR / "figure_3_3_data_relation.png"
    img.save(path)
    assets["relation"] = path

    # Figure 4.1
    img = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(img)
    flow = [
        ((60, 260, 360, 430), "访问 /students"),
        ((420, 260, 760, 430), "requireLogin\n校验会话"),
        ((820, 260, 1180, 430), "StudentController\n解析筛选条件与分页参数"),
        ((1240, 260, 1580, 430), "StudentModel\n查询分页与统计数据"),
    ]
    for idx, (xy, text) in enumerate(flow):
        draw_box(draw, xy, text, box_font, fill="#E8F1FB")
        if idx < len(flow) - 1:
            draw_arrow(draw, (xy[2], (xy[1] + xy[3]) / 2), (flow[idx + 1][0][0], (xy[1] + xy[3]) / 2))
    draw_box(draw, (1240, 520, 1580, 680), "EJS 渲染总览页面\n展示列表、统计与操作入口", box_font, fill="#EEF7EE")
    draw_arrow(draw, (1410, 430), (1410, 520))
    draw_box(draw, (480, 520, 820, 680), "未登录时重定向\n返回 /login", box_font, fill="#FFF2F2")
    draw_arrow(draw, (590, 430), (650, 520))
    draw.text((640, 110), "学生总览访问与数据处理流程", font=title_font, fill="#111111")
    path = ASSET_DIR / "figure_4_1_student_flow.jpg"
    img.save(path, quality=95)
    assets["student_flow"] = path

    # Figure 4.2
    img = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(img)
    stages = [
        ((80, 320, 360, 500), "需求设计\nCodex 辅助原型与代码建议"),
        ((420, 320, 700, 500), "本地开发\n实现功能与修复问题"),
        ((760, 320, 1040, 500), "测试验证\nJest + Supertest"),
        ((1100, 320, 1380, 500), "持续集成\nGitHub Actions"),
        ((1440, 320, 1720, 500), "容器化部署\nDocker / Compose"),
    ]
    for idx, (xy, text) in enumerate(stages):
        draw_box(draw, xy, text, box_font, fill="#E8F1FB" if idx % 2 == 0 else "#F8F5FF")
        if idx < len(stages) - 1:
            draw_arrow(draw, (xy[2], (xy[1] + xy[3]) / 2), (stages[idx + 1][0][0], (xy[1] + xy[3]) / 2))
    draw.text((660, 140), "Codex 与 DevOps 闭环实践路径", font=title_font, fill="#111111")
    path = ASSET_DIR / "figure_4_2_devops_flow.png"
    img.save(path)
    assets["devops"] = path
    return assets


def add_figure(doc, image_path, caption, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_figure_caption(doc, caption)


def add_toc_page(doc):
    p = doc.add_paragraph()
    write_paragraph(p, "目 录", size=Pt(22), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=18, line=22)
    toc = doc.add_paragraph()
    toc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()


def add_cover_page(doc):
    section = doc.sections[0]
    set_page_layout(section)
    section.different_first_page_header_footer = True

    top_para = doc.add_paragraph()
    top_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(top_para, before=0, after=8, line=16)
    if LOGO_PATH.exists():
        top_para.add_run().add_picture(str(LOGO_PATH), width=Cm(7.6))

    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    write_paragraph(title, "计算机学院专业综合设计报告", size=Pt(26), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=40, line=28)
    project = doc.add_paragraph()
    write_paragraph(project, "基于 Codex 与 DevOps 闭环的学生信息管理系统自动化编程流程探索", size=Pt(24), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line=28)
    subtitle = doc.add_paragraph()
    write_paragraph(subtitle, "—— 以 Node.js、Express、MySQL 为核心的工程实践", size=Pt(20), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=50, line=24)

    info_table = doc.add_table(rows=7, cols=2)
    remove_table_borders(info_table)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_geometry.apply_table_geometry(
        info_table,
        table_geometry.exact_column_widths([2200, 5600], total_width_dxa=7800),
        table_width_dxa=7800,
        indent_dxa=780,
    )
    labels = ["姓    名", "班    级", "学    号", "学科专业", "同组成员", "指导教师", "日    期"]
    values = ["（待填写）", "（待填写）", "（待填写）", "计算机科学与技术", "无", "冯国朋", "2026 年 7 月 1 日"]
    for idx, (label, value) in enumerate(zip(labels, values)):
        fill_table_cell(info_table.cell(idx, 0), label, size=Pt(16), align=WD_ALIGN_PARAGRAPH.LEFT)
        fill_table_cell(info_table.cell(idx, 1), value, size=Pt(16), align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_border(info_table.cell(idx, 1), bottom={"val": "single", "sz": "6", "color": "000000"})

    doc.add_paragraph()
    for _ in range(4):
        doc.add_paragraph()
    school = doc.add_paragraph()
    write_paragraph(school, "中原工学院计算机学院", size=Pt(24), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=26)

    next_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_layout(next_section)
    restart_page_numbering(next_section, 1)
    add_footer_page_number(next_section)
    return next_section


def add_reference_paragraph(doc, text):
    p = doc.add_paragraph()
    write_paragraph(p, text, size=CAPTION_SIZE, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=BODY_LINE_PT)


def add_code_listing(doc, title, content):
    add_heading(doc, title, 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(p, before=0, after=10, line=14)
    run = p.add_run(content)
    set_run_font(run, size=Pt(9), bold=False, east_asia="Consolas", ascii_font="Consolas")


def read_file_text(path):
    return Path(path).read_text(encoding="utf-8")


def build_report(output_path):
    ensure_dirs()
    assets = generate_diagrams()

    doc = Document()
    configure_styles(doc)
    add_cover_page(doc)
    add_toc_page(doc)

    add_heading(doc, "1 序言", 1)
    add_heading(doc, "1.1 选题背景", 2)
    add_body_paragraph(
        doc,
        "随着高校学生规模扩大，传统依赖纸质台账或零散电子表格进行学生信息维护的方式逐渐暴露出查询效率低、统计口径不统一、数据更新不同步等问题。"
        "与此同时，生成式人工智能工具正在改变软件开发方式，尤其是在原型搭建、代码补全、缺陷修复和测试生成等环节中，AI 已经能够显著提升开发效率。"
    )
    add_body_paragraph(
        doc,
        "本课题以学生信息管理系统为具体实验载体，将 Codex 辅助编程能力与 GitHub Actions、Docker 等 DevOps 工具结合，尝试打通“需求分析—系统设计—编码实现—测试验证—持续集成—容器化部署”的闭环流程。"
        "项目既满足高校教务场景下的业务需求，又能够体现自动化编程与工程化交付的综合训练目标。"
    )

    add_heading(doc, "1.2 选题目的及意义", 2)
    add_body_paragraph(
        doc,
        "本课题的直接目标是设计并实现一个具备登录认证、学生档案维护、班级管理、统计展示和数据导出能力的学生信息管理系统；"
        "更深层的目标是探索 Codex 在软件开发全过程中的辅助价值，并验证其与测试、持续集成、容器化部署结合后的工程实践效果。"
    )
    add_body_paragraph(
        doc,
        "从实践意义看，本项目一方面能够为学生管理类系统的快速开发提供一个可运行的样例，另一方面能够帮助开发者理解如何将 AI 辅助编码能力嵌入标准 DevOps 流程之中，减少重复劳动，提高迭代速度，并在一定程度上提升系统交付质量。"
    )

    add_heading(doc, "1.3 本人的工作", 2)
    add_body_paragraph(
        doc,
        "在本次综合设计中，本人独立完成了项目选题、需求分析、数据库设计、前后端编码、测试用例编写、GitHub Actions 工作流配置、Docker 容器化部署方案编制以及实验报告撰写等工作。"
    )
    add_body_paragraph(
        doc,
        "具体而言，在开发过程中本人借助 Codex 进行目录结构规划、控制器与模型层代码补全、界面样式优化、单元测试样例生成和配置文件修复，同时对生成结果进行人工审查与迭代修改，确保系统实现与实验要求保持一致。"
    )

    add_heading(doc, "2 课题需求分析", 1)
    add_heading(doc, "2.1 课题功能需求分析", 2)
    add_body_paragraph(
        doc,
        "围绕学生日常档案维护与班级协同管理场景，系统需要支持基础身份校验、学生数据增删改查、班级数据维护、统计展示以及 CSV 数据导出等功能。"
        "为了便于后续实现，本项目将功能需求划分为五个模块。"
    )

    add_heading(doc, "2.1.1 登录认证模块", 3)
    add_body_paragraph(
        doc,
        "系统应提供管理员登录入口，对用户名与密码进行校验；登录成功后建立会话，并限制未登录用户直接访问学生与班级相关页面；管理员可通过退出登录主动销毁会话。"
    )
    add_heading(doc, "2.1.2 学生信息管理模块", 3)
    add_body_paragraph(
        doc,
        "系统应支持学生信息的新增、编辑、删除、分页查询、关键字检索、按性别/班级/状态筛选、详情查看与 CSV 导出。学生实体除学号、姓名、性别、班级等核心字段外，还需维护邮箱、宿舍、紧急联系人、学籍状态等扩展信息。"
    )
    add_heading(doc, "2.1.3 班级管理模块", 3)
    add_body_paragraph(
        doc,
        "系统应支持班级基本信息维护，包括班级名称、年级、专业、辅导员、班长、容量等字段；同时需要对班级人数、在读人数和容量占用情况进行汇总展示。修改班级名称时，应同步更新关联学生的班级字段；删除班级前，需要校验是否仍存在关联学生。"
    )
    add_heading(doc, "2.1.4 数据导出与统计模块", 3)
    add_body_paragraph(
        doc,
        "系统应提供总人数、性别分布、班级分布、状态分布、最近录入学生等统计信息，并支持依据当前筛选条件导出 CSV 文件，满足教务汇总与留档需求。"
    )
    add_heading(doc, "2.1.5 自动化工程化模块", 3)
    add_body_paragraph(
        doc,
        "系统除业务功能外，还应具备基础工程化能力，包括测试用例执行、持续集成工作流校验和 Docker 容器化部署配置，以支撑后续的稳定交付和迭代维护。"
    )

    add_data_table(
        doc,
        ["功能模块", "核心需求", "输入/输出", "实现要点"],
        [
            ["登录认证", "管理员登录、会话控制、退出登录", "输入：账号密码；输出：登录状态", "bcrypt 校验密码，session 保存用户信息"],
            ["学生管理", "新增、编辑、删除、查询、详情", "输入：学生表单；输出：学生列表/详情页", "控制器处理校验，模型执行 CRUD"],
            ["班级管理", "班级维护、人数统计、容量管理", "输入：班级表单；输出：班级面板", "支持班级更名联动学生数据"],
            ["导出统计", "统计展示与 CSV 导出", "输入：筛选条件；输出：统计结果/CSV", "支持分页、状态统计、最近录入"],
            ["工程化", "测试、CI、容器化", "输入：代码提交；输出：测试结果/镜像", "Jest、GitHub Actions、Docker 协同"],
        ],
        [2.2, 4.3, 2.0, 4.5],
        "表2.1 系统功能需求汇总",
    )

    add_heading(doc, "2.2 课题性能需求分析", 2)
    add_body_paragraph(
        doc,
        "由于本系统主要面向课程实验与中小规模教务维护场景，因此性能需求侧重于常见查询、筛选和表单提交过程的及时响应，以及后续扩展时的可维护性。系统在功能设计上采用分页查询与条件拼接方式，避免一次性渲染过多数据。"
    )
    add_data_table(
        doc,
        ["指标项", "需求说明", "项目中的应对策略"],
        [
            ["响应时间", "常见页面请求应保持秒级返回", "学生列表采用分页查询，统计结果按需计算"],
            ["并发适应性", "满足课程实验与小规模管理并发访问", "Express + MySQL 连接池满足轻量场景"],
            ["数据一致性", "班级与学生关联信息保持同步", "班级更名时同步更新 students.class_name"],
            ["可维护性", "后续便于补列、扩展模块", "schemaSync 自动建表与补列，路由控制器分层"],
            ["可测试性", "关键路径可通过自动化测试验证", "Jest + Supertest 覆盖主要访问流程"],
        ],
        [2.0, 4.0, 4.65],
        "表2.2 性能与质量需求分析",
    )

    add_heading(doc, "2.3 其他需求", 2)
    add_body_paragraph(
        doc,
        "除功能与性能需求外，系统还需满足以下附加要求：一是安全性要求，未登录用户不能访问核心业务页面，敏感操作应由管理员执行；二是可用性要求，界面布局应清晰、表单提示应直观，便于教务人员快速上手；三是可部署性要求，系统需提供标准化的 Dockerfile 与 docker-compose.yml，便于在不同环境中快速启动。"
    )

    add_heading(doc, "3 系统概要设计", 1)
    add_heading(doc, "3.1 系统体系结构", 2)
    add_body_paragraph(
        doc,
        "系统采用典型的 B/S 三层组织方式，并结合工程化支撑层形成面向实验实践的轻量架构。浏览器端通过 EJS 模板和 Bootstrap 页面与用户交互，控制器层负责请求解析、参数校验与业务编排，模型层负责与 MySQL 数据库交互，Codex、测试框架、持续集成和 Docker 则作为工程化支撑能力贯穿开发与交付全过程。"
    )
    add_figure(doc, assets["architecture"], "图3.1 系统总体体系结构图", width_cm=14.6)

    add_heading(doc, "3.2 系统功能设计", 2)
    add_body_paragraph(
        doc,
        "依据需求分析结果，系统功能可划分为登录认证、学生管理、班级管理、数据导出与统计展示、工程化支撑五个部分。各功能之间关系清晰：登录认证作为访问前置条件，学生管理与班级管理共同构成核心业务域，统计与导出建立在学生数据集合之上，而工程化能力保障整体开发与交付过程。"
    )
    add_figure(doc, assets["function"], "图3.2 系统功能结构图", width_cm=14.6)
    add_data_table(
        doc,
        ["模块", "主要页面或接口", "职责说明"],
        [
            ["认证模块", "/login, /logout", "完成管理员登录校验、会话写入与注销"],
            ["学生模块", "/students 及其子路由", "完成学生数据维护、条件查询、详情展示与导出"],
            ["班级模块", "/classes 及其子路由", "完成班级资料维护、人数统计与容量展示"],
            ["统计模块", "学生总览页、班级面板", "展示性别、状态、班级、最近录入等统计结果"],
            ["工程化模块", "npm test / GitHub Actions / Docker", "保障开发质量与部署一致性"],
        ],
        [2.1, 4.3, 5.4],
        "表3.1 系统模块职责划分",
    )

    add_heading(doc, "3.3 主要数据结构概要设计", 2)
    add_body_paragraph(
        doc,
        "系统采用面向业务对象的数据组织方式，核心数据结构包括用户、学生、班级三类。用户结构负责登录认证，学生结构负责承载教务档案信息，班级结构负责描述班级维度的组织信息。"
    )
    add_heading(doc, "3.3.1 学生数据结构", 3)
    add_data_table(
        doc,
        ["字段名", "类型", "说明"],
        [
            ["student_no", "VARCHAR(30)", "学生学号，唯一标识学生记录"],
            ["name", "VARCHAR(50)", "学生姓名"],
            ["gender", "VARCHAR(10)", "性别"],
            ["class_name", "VARCHAR(60)", "所属班级名称"],
            ["phone / email", "VARCHAR", "联系方式"],
            ["birthday", "DATE", "出生日期"],
            ["status", "VARCHAR(20)", "在读、实习、休学等状态"],
            ["dormitory", "VARCHAR(50)", "宿舍信息"],
            ["emergency_contact / emergency_phone", "VARCHAR", "紧急联系人与电话"],
            ["notes", "TEXT", "备注信息"],
        ],
        [2.6, 2.4, 4.35],
        "表3.2 学生数据结构概要说明",
    )
    add_heading(doc, "3.3.2 班级数据结构", 3)
    add_data_table(
        doc,
        ["字段名", "类型", "说明"],
        [
            ["class_name", "VARCHAR(60)", "班级名称，唯一标识班级"],
            ["grade", "VARCHAR(20)", "年级"],
            ["major", "VARCHAR(50)", "所属专业"],
            ["counselor", "VARCHAR(50)", "辅导员姓名"],
            ["monitor_name", "VARCHAR(50)", "班长姓名"],
            ["student_capacity", "INT", "班级容量"],
            ["created_at", "TIMESTAMP", "创建时间"],
        ],
        [2.6, 2.4, 4.35],
        "表3.3 班级数据结构概要说明",
    )

    add_heading(doc, "3.4 主要数据库表概要设计", 2)
    add_body_paragraph(
        doc,
        "数据库层采用 MySQL 关系型存储，围绕 users、classes、students 三张核心表展开。其中 students.class_name 与 classes.class_name 形成逻辑关联，用于实现班级维度的统计与学生数据联动。为提高初始化效率，系统额外设计了 schemaSync 逻辑，用于在启动时自动建表、补充缺失字段并同步基础账号数据。"
    )
    add_figure(doc, assets["relation"], "图3.3 核心数据关系示意图", width_cm=14.2)
    add_data_table(
        doc,
        ["表名", "主要字段", "作用说明"],
        [
            ["users", "id, username, password", "保存管理员登录账号与加密密码"],
            ["classes", "class_name, grade, major, counselor, student_capacity", "保存班级基础信息与容量配置"],
            ["students", "student_no, name, class_name, status, phone, notes", "保存学生档案与状态信息"],
        ],
        [2.2, 5.0, 2.8],
        "表3.4 主要数据库表设计概览",
    )

    add_heading(doc, "3.5 主要接口概要设计", 2)
    add_body_paragraph(
        doc,
        "系统接口采用基于 Express Router 的资源组织方式。认证接口承担登录与注销功能，学生与班级接口分别负责业务数据的查询和维护。由于项目采用服务端渲染模式，接口输出既包括 HTML 页面，也包括用于导出的 CSV 数据。"
    )
    add_data_table(
        doc,
        ["接口/路由", "方法", "功能说明"],
        [
            ["/login", "GET / POST", "展示登录页并完成登录认证"],
            ["/logout", "POST", "注销当前会话并返回登录页"],
            ["/students", "GET / POST", "学生列表查询与新增学生"],
            ["/students/:id", "GET", "查看学生详情"],
            ["/students/:id/update", "POST", "更新学生信息"],
            ["/students/:id/delete", "POST", "删除学生记录"],
            ["/students/export", "GET", "导出符合筛选条件的 CSV"],
            ["/classes", "GET / POST", "班级面板展示与新增班级"],
            ["/classes/:id/update", "POST", "更新班级信息并同步学生班级名称"],
            ["/classes/:id/delete", "POST", "删除空班级"],
        ],
        [4.0, 1.8, 3.2],
        "表3.5 主要接口设计说明",
    )

    add_heading(doc, "4 系统详细设计与实现", 1)
    add_heading(doc, "4.1 登录认证子系统的设计与实现", 2)
    add_body_paragraph(
        doc,
        "登录认证子系统主要由 authRoutes、authController、userModel 和 authMiddleware 组成。用户首先访问 /login 页面输入账号密码，控制器使用 userModel 查询数据库中的管理员记录，再通过 bcrypt.compare 对明文密码和哈希值进行比对。"
    )
    add_body_paragraph(
        doc,
        "认证成功后，系统将用户 id 和 username 写入 session，并统一重定向到学生总览页；若认证失败，则在当前页面给出错误提示。为了保证业务安全性，studentRoutes 与 classRoutes 中的所有核心接口均接入 requireLogin 中间件，未登录时直接跳转回登录页。"
    )

    add_heading(doc, "4.2 学生信息管理子系统的设计与实现", 2)
    add_body_paragraph(
        doc,
        "学生信息管理子系统是本项目的核心部分。StudentController 负责解析筛选条件、分页参数和表单数据；StudentModel 则负责拼接 SQL 条件，实现分页查询、统计汇总、详情查询、新增、修改、删除与导出等操作。"
    )
    add_body_paragraph(
        doc,
        "在学生总览页中，系统通过 Promise.all 并行查询分页数据、性别统计、班级统计、状态统计和最近录入学生列表，提高页面装载效率；在新增和编辑过程中，系统会先执行字段非空校验，再进行学号唯一性校验，最后调用模型层完成数据库写入。CSV 导出功能则根据当前过滤条件调用 findAll，将结果拼接为带 BOM 头的文本响应，方便在 Excel 中直接打开。"
    )

    add_heading(doc, "4.3 班级管理子系统的设计与实现", 2)
    add_body_paragraph(
        doc,
        "班级管理子系统用于维护班级的组织信息，并为学生管理页面提供班级数据来源。ClassModel 通过 LEFT JOIN students 的方式统计班级人数与在读人数，ClassController 则负责班级表单校验、重复班级名称检测、班级更名联动以及删除前约束校验。"
    )
    add_body_paragraph(
        doc,
        "当管理员修改班级名称时，系统会先更新 classes 表，再调用 renameStudents 将 students 表中对应的 class_name 一并修改，从而保证班级与学生之间的数据一致性；删除班级时，则通过 countStudentsByName 判断班级下是否仍有学生，防止产生无归属学生记录。"
    )

    add_heading(doc, "4.4 主要逻辑流程/关键实现设计与实现", 2)
    add_heading(doc, "4.4.1 主要系统逻辑流程", 3)
    add_body_paragraph(
        doc,
        "系统的核心访问路径主要发生在学生总览页面。用户经过认证后访问 /students，控制器统一处理筛选条件和分页参数，模型层根据条件动态拼接 WHERE 子句并执行分页查询与统计查询，最终由 EJS 模板将学生列表、统计卡片和辅助模块整合渲染到页面中。"
    )
    add_heading(doc, "4.4.2 数据库初始化与自同步实现", 3)
    add_body_paragraph(
        doc,
        "考虑到实验项目在不同机器和容器环境中运行时可能存在数据库表缺失或字段不一致的问题，项目设计了 schemaSync 机制。应用启动时会先执行 ensureSchema：依次检查 users、classes、students 三张表是否存在，不存在时自动建表；若表已存在，则继续检查 email、birthday、status 等扩展字段是否缺失，并通过 ALTER TABLE 方式补列。"
    )
    add_body_paragraph(
        doc,
        "此外，schemaSync 还会自动写入管理员默认账号，并根据 students 表中已有的 class_name 自动补种 classes 数据，从而提高系统首次部署和后续迁移时的可用性。这一设计简化了实验环境搭建流程，也体现了本项目对自动化维护能力的关注。"
    )

    add_heading(doc, "4.4.3 Codex 辅助开发过程", 3)
    add_body_paragraph(
        doc,
        "Codex 在本项目中主要承担了四类工作：第一，围绕学生管理场景辅助拆解功能模块和目录结构；第二，针对控制器、模型和视图层重复性较强的代码进行高效补全；第三，在页面样式调整、路由扩展、CSV 导出、测试桩编写等环节给出可直接落地的实现建议；第四，在发现编码或逻辑错误后，通过多轮提示进行修复与重构。"
    )
    add_body_paragraph(
        doc,
        "在实际使用过程中，Codex 的价值不仅体现在“生成代码”，更体现在“加速迭代”。开发者可以先给出目标、边界和已有代码上下文，再根据生成结果进行人工判断和修改，形成“人制定约束，AI 提供候选实现，人再校正”的协作方式。"
    )

    add_heading(doc, "4.5 测试与结果分析", 2)
    add_body_paragraph(
        doc,
        "项目使用 Jest 和 Supertest 对关键路由进行自动化测试，重点覆盖健康检查、登录页面加载、未登录重定向、学生总览页渲染、CSV 导出和班级面板渲染等核心路径。执行结果表明，当前测试集共包含 1 个测试套件、6 个测试用例，均通过验证。"
    )
    add_data_table(
        doc,
        ["测试编号", "测试内容", "预期结果", "实际结果"],
        [
            ["T1", "GET /health 健康检查", "返回 200 且内容为 ok", "通过"],
            ["T2", "GET /login 页面访问", "正常渲染登录页", "通过"],
            ["T3", "未登录访问 /students", "302 重定向到 /login", "通过"],
            ["T4", "登录后访问 /students", "渲染学生总览及统计信息", "通过"],
            ["T5", "登录后访问 /students/export", "成功下载 CSV 内容", "通过"],
            ["T6", "登录后访问 /classes", "渲染班级管理面板", "通过"],
        ],
        [1.6, 4.2, 3.0, 1.2],
        "表4.1 核心测试用例与执行结果",
    )
    add_body_paragraph(
        doc,
        "从结果上看，现有测试已经覆盖系统的主要访问入口和关键业务链路，能够在后续迭代中及时发现登录跳转、页面渲染和导出逻辑的回归问题。由于当前项目尚未引入真实数据库集成测试与前端自动化浏览器测试，因此后续仍可继续扩展测试深度。"
    )

    add_heading(doc, "4.6 GitHub Actions 与 Docker 工程化实现", 2)
    add_body_paragraph(
        doc,
        "为了验证“代码提交—测试—镜像构建”的自动化闭环，本项目在 .github/workflows/ci.yml 中定义了 CI 工作流：当代码推送到 main 分支或发起 Pull Request 时，GitHub Actions 会自动检出代码、配置 Node.js 20 环境、执行 npm ci 安装依赖、运行 npm test 执行测试，最后执行 docker build 构建应用镜像。"
    )
    add_body_paragraph(
        doc,
        "在部署层面，项目通过 Dockerfile 选择 node:20-alpine 作为基础镜像，并在 docker-compose.yml 中配置应用服务端口、数据库连接参数和 SESSION_SECRET 等环境变量。Compose 方案通过 host.docker.internal 访问主机 MySQL 服务，使本地实验环境下的 Node.js 容器能够快速接入数据库。"
    )
    add_figure(doc, assets["devops"], "图4.2 Codex 与 DevOps 闭环流程图", width_cm=14.6)
    add_data_table(
        doc,
        ["工程化项", "配置位置", "作用说明"],
        [
            ["持续集成", ".github/workflows/ci.yml", "自动执行依赖安装、测试和镜像构建"],
            ["镜像构建", "Dockerfile", "定义 Node.js 运行环境与容器启动方式"],
            ["本地编排", "docker-compose.yml", "统一应用端口、数据库地址和运行参数"],
            ["启动校验", "app.js + schemaSync", "在服务监听前确保数据库结构可用"],
        ],
        [2.1, 3.5, 3.75],
        "表4.2 工程化实现要点汇总",
    )

    add_heading(doc, "5 实训总结", 1)
    add_heading(doc, "5.1 技术总结", 2)
    add_body_paragraph(
        doc,
        "通过本次综合设计，本人系统掌握了基于 Node.js、Express、EJS 与 MySQL 构建中小型 Web 管理系统的基本流程，进一步理解了 MVC 分层、数据库建模、会话认证、分页查询、统计展示、CSV 导出、自动化测试和容器化部署等关键技术点。"
    )
    add_body_paragraph(
        doc,
        "项目开发过程也证明，Codex 与标准工程化工具并不是彼此孤立的。将 Codex 用于需求拆解、代码补全和问题修复，将 GitHub Actions 用于自动校验，将 Docker 用于标准化运行环境，可以有效缩短从需求到交付的路径，提高实验项目的完整度。"
    )

    add_heading(doc, "5.2 思想总结", 2)
    add_body_paragraph(
        doc,
        "本次实训不仅是一次编程任务，更是一次对软件工程方法的实践训练。本人在开发中深刻体会到：功能实现只是起点，真正决定项目质量的往往是需求是否清晰、结构是否合理、测试是否到位、交付是否标准化。"
    )
    add_body_paragraph(
        doc,
        "同时，AI 工具的加入并没有削弱开发者的作用，反而更要求开发者具备明确目标、审查结果和持续修正的能力。今后在继续学习和项目实践中，本人将进一步提升需求表达、系统设计、测试设计和工程化交付能力，使 AI 真正成为可靠的开发协作伙伴。"
    )

    add_heading(doc, "参考文献", 1)
    references = [
        "[1] 萨师煊, 王珊. 数据库系统概论[M]. 北京: 高等教育出版社, 2014.",
        "[2] Node.js. Node.js Documentation[EB/OL]. https://nodejs.org/.",
        "[3] Express.js. Express Web Framework Documentation[EB/OL]. https://expressjs.com/.",
        "[4] Oracle. MySQL Reference Manual[EB/OL]. https://dev.mysql.com/doc/.",
        "[5] Docker Inc. Docker Documentation[EB/OL]. https://docs.docker.com/.",
        "[6] GitHub. GitHub Actions Documentation[EB/OL]. https://docs.github.com/actions.",
        "[7] OpenAI. Codex Documentation[EB/OL]. https://platform.openai.com/docs/.",
    ]
    for item in references:
        add_reference_paragraph(doc, item)

    add_heading(doc, "附录", 1)
    add_body_paragraph(
        doc,
        "附录部分给出项目中具有代表性的工程化配置片段，以说明系统在持续集成、容器化和数据库自同步方面的实现方式。"
    )

    ci_text = read_file_text(ROOT / ".github/workflows/ci.yml")
    compose_text = read_file_text(ROOT / "docker-compose.yml")
    schema_text = read_file_text(ROOT / "config/schemaSync.js")

    add_code_listing(doc, "附录A GitHub Actions 工作流关键配置", ci_text.strip())
    add_code_listing(doc, "附录B Docker Compose 关键配置", compose_text.strip())
    add_code_listing(doc, "附录C schemaSync 核心逻辑摘录", "\n".join(schema_text.strip().splitlines()[:80]))

    doc.save(str(output_path))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", default=str(OUTPUT_DIR / "专业综合设计报告_学生信息管理系统.docx"))
    args = parser.parse_args()
    build_report(Path(args.output))


if __name__ == "__main__":
    main()
