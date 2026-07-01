import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
import pypdfium2 as pdfium


def inspect_docx(path: Path) -> dict:
    doc = Document(str(path))
    data = {
        "paragraphs": [],
        "tables": []
    }

    for index, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if text:
            data["paragraphs"].append({"index": index, "text": text})

    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row_index, row in enumerate(table.rows, start=1):
            rows.append(
                {
                    "index": row_index,
                    "cells": [cell.text.replace("\n", " ").strip() for cell in row.cells]
                }
            )
        data["tables"].append(
            {
                "index": table_index,
                "row_count": len(table.rows),
                "col_count": len(table.columns),
                "rows": rows
            }
        )

    return data


def duplicate_template(template_path: Path, output_path: Path) -> None:
    doc = Document(str(template_path))
    doc.save(str(output_path))


def render_pdf(pdf_path: Path, output_dir: Path, scale: float) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf = pdfium.PdfDocument(str(pdf_path))

    for index in range(len(pdf)):
        page = pdf[index]
        image = page.render(scale=scale).to_pil()
        image.save(output_dir / f"page-{index + 1}.png")


def set_run_font(run, size=12, bold=False):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold


def reset_cell(cell):
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    cell.add_paragraph()


def fill_multiline_cell(cell, lines, size=12, spacing=20):
    reset_cell(cell)
    first = True
    for line in lines:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.35
        run = paragraph.add_run(line)
        set_run_font(run, size=size)


def fill_plain_cell(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    reset_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_paragraph_text(paragraph, text, size=12, bold=False):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def generate_weekly_report(template_path: Path, output_path: Path) -> None:
    weeks = [
        {
            "date": "2026 年 6 月 15 日-------2026 年 6 月 21 日",
            "completed": [
                "1. 围绕《基于 Codex 与 DevOps 闭环的自动化编程流程探索》明确实训目标，梳理“需求分析-系统设计-编码实现-测试部署”的整体路线。",
                "2. 完成学生信息管理系统选题和技术方案确定，选用 Node.js、Express、EJS、MySQL 作为主要开发栈，规划登录认证、学生管理、班级管理等核心模块。",
                "3. 借助 Codex 辅助完成项目目录分层设计，初步搭建 controllers、models、routes、views、config 等基础结构，并整理实验报告提纲与提示词资料。",
                "4. 搭建本地开发环境，初始化 npm 工程和依赖包，调研数据库连接、容器化部署、持续集成所需的关键配置项。",
                "5. 收集 Codex 辅助编程、GitHub Actions、Docker 与自动化测试相关资料，为后续形成 DevOps 闭环做准备，整体完成情况符合周计划。",
            ],
            "next": [
                "1. 完成系统基础框架编码，打通登录、学生信息录入、查询与编辑的主流程。",
                "2. 设计 students、classes、users 等数据表并编写初始化 SQL 脚本。",
                "3. 继续细化页面原型和表单交互，形成可运行的第一版系统原型。",
            ],
            "progress": "较快\n正常（√）\n较慢",
            "comment": "本周已完成前期调研与方案设计，技术路线较清晰，能够结合实验要求主动整理资料并明确任务分工。建议下周尽快落实核心业务功能，实现从方案到代码的阶段性闭环。"
        },
        {
            "date": "2026 年 6 月 22 日-------2026 年 6 月 28 日",
            "completed": [
                "1. 完成 Express 应用主体搭建，配置 session、method-override、静态资源和健康检查接口，系统可以正常启动运行。",
                "2. 实现登录认证以及学生信息模块的新增、编辑、删除、分页、关键字筛选、状态筛选、详情查看和 CSV 导出功能。",
                "3. 完成 MySQL 数据模型与 init.sql 编写，补充邮箱、宿舍、紧急联系人、学籍状态等字段，并增加 schemaSync 逻辑支持自动建表和补列。",
                "4. 借助 Codex 辅助优化前端页面，完成登录页、学生总览页、学生详情页和录入表单的 EJS 模板及样式设计。",
                "5. 编写 Jest + Supertest 测试用例，对健康检查、登录跳转、学生列表渲染、导出接口等关键流程进行验证，形成“编码-测试-修正”的迭代闭环。",
            ],
            "next": [
                "1. 新增班级管理模块，完善班级与学生数据之间的联动关系。",
                "2. 补充 Dockerfile、docker-compose.yml 和 GitHub Actions 配置，验证持续集成与容器化流程。",
                "3. 继续完善异常处理、测试覆盖和实验记录，为总结 Codex 辅助开发效果积累材料。",
            ],
            "progress": "较快（√）\n正常\n较慢",
            "comment": "本周核心功能推进较快，系统原型已经具备基本可用性。能够结合 Codex 的代码补全与问题修复能力开展开发，建议下周继续加强模块联动和自动化环节的验证。"
        },
        {
            "date": "2026 年 6 月 29 日-------2026 年 7 月 5 日",
            "completed": [
                "1. 在现有学生管理基础上新增班级管理模块，实现班级新增、编辑、删除、容量统计、班级人数汇总以及班级更名时学生数据同步更新。",
                "2. 抽离侧边栏等公共视图，补充班级管理面板、学生详情页等界面，进一步优化系统整体信息展示和操作连贯性。",
                "3. 编写 GitHub Actions CI 工作流，完成 npm ci、npm test 和 docker build 的自动化执行，初步实现代码提交后的持续集成校验。",
                "4. 完成 Dockerfile 与 docker-compose.yml 配置，打通 Node.js 应用与 MySQL 相关运行参数，形成可复用的容器化部署方案。",
                "5. 持续补充模拟数据和测试桩，验证登录、学生模块、班级模块的主要访问路径；同时完善实验报告提纲，准备总结 Codex 与 DevOps 闭环实践结果。",
                "6. 阶段性完成“代码实现-测试验证-持续集成-容器化部署”的实验主线，当前整体进度较为顺利。",
            ],
            "next": [
                "1. 完成实验报告撰写，整理系统截图、测试结果和关键配置说明。",
                "2. 补充异常场景测试，继续优化中文编码、表单校验和部署细节。",
                "3. 总结 Codex 在原型设计、代码补全、缺陷修复和 DevOps 联调中的作用与不足，完善结题材料。",
            ],
            "progress": "较快（√）\n正常\n较慢",
            "comment": "本周任务完成较全面，已初步形成 Codex 与 DevOps 闭环的实践成果。后续应重点做好实验总结与结果量化，进一步提升系统稳定性和文档完整性。"
        }
    ]

    doc = Document(str(template_path))

    date_paragraph_indices = [1, 4, 8]
    for paragraph_index, week in zip(date_paragraph_indices, weeks):
        set_paragraph_text(doc.paragraphs[paragraph_index], week["date"], size=12)

    for table, week in zip(doc.tables, weeks):
        fill_multiline_cell(table.rows[1].cells[0], week["completed"], size=12)
        fill_multiline_cell(table.rows[3].cells[0], week["next"], size=12)

        footer_cells = table.rows[4].cells
        if len(footer_cells) >= 3:
            fill_plain_cell(footer_cells[2], week["progress"], size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        if len(footer_cells) >= 4:
            fill_multiline_cell(footer_cells[3], [week["comment"]], size=11)
        elif len(footer_cells) >= 3:
            fill_multiline_cell(footer_cells[-1], [week["comment"]], size=11)

    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    subparsers = parser.add_subparsers(dest="command", required=True)

    inspect_parser = subparsers.add_parser("inspect")
    inspect_parser.add_argument("path")

    copy_parser = subparsers.add_parser("copy")
    copy_parser.add_argument("template")
    copy_parser.add_argument("output")

    render_parser = subparsers.add_parser("render-pdf")
    render_parser.add_argument("pdf")
    render_parser.add_argument("output_dir")
    render_parser.add_argument("--scale", type=float, default=2.0)

    generate_parser = subparsers.add_parser("generate")
    generate_parser.add_argument("template")
    generate_parser.add_argument("output")

    args = parser.parse_args()

    if args.command == "inspect":
        info = inspect_docx(Path(args.path))
        print(json.dumps(info, ensure_ascii=False, indent=2))
        return

    if args.command == "copy":
        duplicate_template(Path(args.template), Path(args.output))
        return

    if args.command == "render-pdf":
        render_pdf(Path(args.pdf), Path(args.output_dir), args.scale)
        return

    if args.command == "generate":
        generate_weekly_report(Path(args.template), Path(args.output))


if __name__ == "__main__":
    main()
